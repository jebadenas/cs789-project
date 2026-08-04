"""Ingest Canvas reflective-journal exports into a pseudonymised table.

Walks the untouched Canvas exports under ``data/journals/raw/``, parses the
Canvas filename convention, extracts plain text (PDF + DOCX), pseudonymises the
student identity, and writes two artefacts to ``data/journals/processed/``:

* ``entries.parquet``      — one row per journal file, including extracted text.
* ``entry_manifest.csv``   — the same table minus ``text``, so there is a
  reviewable manifest that carries no journal content.

plus the identity map to ``data/journals/crosswalk/name_to_anon.csv``.

Everything downstream references only ``anon_id``. Filename pseudonymisation is
NOT text de-identification — entry bodies name teammates, so anything quoted
verbatim in §5.4 must be manually scrubbed first.

Nothing this module writes is ever committed (``data/`` is git-ignored).

Run:  ``python3 -m src.qualitative.ingest``
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import logging
import re
import warnings
from dataclasses import asdict, dataclass, field
from datetime import datetime
from pathlib import Path

import pandas as pd

logger = logging.getLogger(__name__)

# --------------------------------------------------------------------------- #
# Paths (repo-relative; all git-ignored)
# --------------------------------------------------------------------------- #
RAW_DIR = Path("data/journals/raw")
PROCESSED_DIR = Path("data/journals/processed")
CROSSWALK_DIR = Path("data/journals/crosswalk")

ENTRIES_PARQUET = PROCESSED_DIR / "entries.parquet"
ENTRY_MANIFEST_CSV = PROCESSED_DIR / "entry_manifest.csv"
UNPARSED_CSV = PROCESSED_DIR / "unparsed_files.csv"
CROSSWALK_CSV = CROSSWALK_DIR / "name_to_anon.csv"

# --------------------------------------------------------------------------- #
# Constants
# --------------------------------------------------------------------------- #
# Canvas bulk-download convention:
#   lastnamefirstname[_LATE]_<canvasUserId>_<submissionId>_<originalFilename>.<ext>
# The name prefix is already lowercased-alphanumeric; the two numeric IDs are
# the reliable anchors. `_LATE` is injected *between* the name and the first ID,
# so we match it explicitly rather than by positional split.
_FILENAME_RE = re.compile(
    r"^(?P<name>.+?)"
    r"_(?:(?P<late>LATE)_)?"
    r"(?P<canvas_id>\d+)_(?P<submission_id>\d+)_"
    r"(?P<orig>.+)$"
)

# Canvas appends -1, -2, ... to the original filename for resubmissions.
_RESUBMISSION_RE = re.compile(r"-(?P<n>\d+)$")

# journal_1 / "journal 1" / journal1 → index
_JOURNAL_IDX_RE = re.compile(r"journal[ _]?(\d+)", re.IGNORECASE)

# Fixed salt: makes ANON_IDs deterministic across runs (so re-ingesting is
# idempotent) while being non-obvious. The crosswalk holds the real mapping;
# the salt is not a security control, just a stable, non-guessable prefix.
_ANON_SALT = "cs789-journals-v1"

# extract_status vocabulary
STATUS_OK = "ok"
STATUS_SUSPECT = "extract_suspect"  # extracted but < MIN_WORDS words
STATUS_EMPTY = "empty"              # supported format, zero text out
STATUS_UNSUPPORTED = "unsupported"  # .doc/.pages/.rtf/.jpg — flagged, not converted
STATUS_ERROR = "error"             # extractor raised

MIN_WORDS = 50  # below this, flag as extract_suspect (likely screenshot-only)

PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
# Present in the export but deliberately not converted — reported, not dropped.
UNSUPPORTED_EXTS = {".doc", ".pages", ".rtf", ".jpg", ".jpeg", ".png", ".md", ".txt"}


# --------------------------------------------------------------------------- #
# Row schema
# --------------------------------------------------------------------------- #
@dataclass
class Entry:
    anon_id: str
    cohort: str
    journal_index: int | None
    canvas_user_id: str
    submission_id: str
    is_late: bool
    is_resubmission: bool
    resubmission_n: int | None
    file_ext: str
    file_bytes: int
    extract_status: str
    word_count: int
    char_count: int
    doc_created: str  # ISO-8601 from document metadata, or "" if unrecoverable
    doc_modified: str
    source_path: str  # identifiable; manifest is git-ignored, kept for review
    text: str = field(default="", repr=False)


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #
def normalise_name(raw: str) -> str:
    """Lowercase and strip all non-alphanumerics — the peer-data join key.

    ``"Bartlett-Wright, Jared"`` and the filename prefix ``bartlettwrightjared``
    both collapse to ``bartlettwrightjared``.
    """
    return re.sub(r"[^a-z0-9]", "", raw.lower())


def anon_id_for(cohort: str, normalised_name: str) -> str:
    """Stable pseudonym for ``(cohort, normalised_name)``."""
    digest = hashlib.sha256(
        f"{_ANON_SALT}:{cohort}:{normalised_name}".encode("utf-8")
    ).hexdigest()
    return "A_" + digest[:12]


def parse_filename(stem: str) -> dict | None:
    """Parse a Canvas filename stem (extension already removed).

    Returns a dict with keys ``normalised_name, is_late, canvas_id,
    submission_id, is_resubmission, resubmission_n`` or ``None`` if the stem
    does not match the convention (caller records it as unparsed).
    """
    m = _FILENAME_RE.match(stem)
    if not m:
        return None

    orig = m.group("orig")
    resub = _RESUBMISSION_RE.search(orig)

    return {
        "normalised_name": normalise_name(m.group("name")),
        "is_late": m.group("late") is not None,
        "canvas_id": m.group("canvas_id"),
        "submission_id": m.group("submission_id"),
        "is_resubmission": resub is not None,
        "resubmission_n": int(resub.group("n")) if resub else None,
    }


def _parse_pdf_date(value: str | None) -> str:
    """Parse a PDF ``D:YYYYMMDDHHmmSS`` date string to an ISO date, else ""."""
    if not value:
        return ""
    m = re.search(r"(\d{4})(\d{2})(\d{2})", value)
    if not m:
        return ""
    try:
        return f"{m.group(1)}-{m.group(2)}-{m.group(3)}"
    except Exception:  # pragma: no cover - defensive
        return ""


def _iso(dt: datetime | None) -> str:
    return dt.date().isoformat() if isinstance(dt, datetime) else ""


def extract_pdf(path: Path) -> tuple[str, str, str]:
    """Return (text, doc_created, doc_modified) for a PDF."""
    import pdfplumber

    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        with pdfplumber.open(path) as pdf:
            pages = [(p.extract_text() or "") for p in pdf.pages]
            meta = pdf.metadata or {}
    text = "\n".join(pages).strip()
    return (
        text,
        _parse_pdf_date(meta.get("CreationDate")),
        _parse_pdf_date(meta.get("ModDate")),
    )


def extract_docx(path: Path) -> tuple[str, str, str]:
    """Return (text, doc_created, doc_modified) for a DOCX."""
    import docx

    document = docx.Document(str(path))
    paras = [p.text for p in document.paragraphs]
    # Include table cell text — reflective journals sometimes use table templates.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paras.append(cell.text)
    text = "\n".join(paras).strip()
    cp = document.core_properties
    return text, _iso(cp.created), _iso(cp.modified)


def _classify(word_count: int, had_error: bool, unsupported: bool) -> str:
    if unsupported:
        return STATUS_UNSUPPORTED
    if had_error:
        return STATUS_ERROR
    if word_count == 0:
        return STATUS_EMPTY
    if word_count < MIN_WORDS:
        return STATUS_SUSPECT
    return STATUS_OK


# --------------------------------------------------------------------------- #
# Main ingest
# --------------------------------------------------------------------------- #
def ingest(raw_dir: Path = RAW_DIR) -> pd.DataFrame:
    """Walk ``raw_dir`` and build the pseudonymised entries table.

    Writes ``entries.parquet``, ``entry_manifest.csv``, ``unparsed_files.csv``
    and the crosswalk. Returns the entries DataFrame (including text).
    """
    if not raw_dir.exists():
        raise FileNotFoundError(f"{raw_dir} does not exist — extract the export first.")

    entries: list[Entry] = []
    unparsed: list[dict] = []
    # (cohort, normalised_name) -> set of canvas_ids seen (collision detection)
    crosswalk: dict[tuple[str, str], set[str]] = {}

    cohorts = sorted(d.name for d in raw_dir.iterdir() if d.is_dir())
    logger.info("Cohorts found: %s", ", ".join(cohorts))

    for cohort in cohorts:
        cohort_dir = raw_dir / cohort
        # Only descend into journal_* folders; skip roster files at cohort root.
        journal_dirs = sorted(
            d for d in cohort_dir.iterdir()
            if d.is_dir() and _JOURNAL_IDX_RE.search(d.name)
        )
        for jdir in journal_dirs:
            jm = _JOURNAL_IDX_RE.search(jdir.name)
            journal_index = int(jm.group(1)) if jm else None

            for f in sorted(jdir.iterdir()):
                if not f.is_file() or f.name.startswith("."):
                    continue
                ext = f.suffix.lower()
                stem = f.name[: -len(f.suffix)] if f.suffix else f.name

                parsed = parse_filename(stem)
                if parsed is None:
                    unparsed.append(
                        {"cohort": cohort, "journal_index": journal_index,
                         "filename": f.name}
                    )
                    continue

                key = (cohort, parsed["normalised_name"])
                crosswalk.setdefault(key, set()).add(parsed["canvas_id"])
                anon = anon_id_for(cohort, parsed["normalised_name"])

                # --- text extraction ---
                text, created, modified = "", "", ""
                had_error = False
                unsupported = ext in UNSUPPORTED_EXTS
                if ext in PDF_EXTS:
                    try:
                        text, created, modified = extract_pdf(f)
                    except Exception as exc:  # noqa: BLE001 - report, don't crash
                        had_error = True
                        logger.warning("PDF extract failed (%s): %s", f.name, exc)
                elif ext in DOCX_EXTS:
                    try:
                        text, created, modified = extract_docx(f)
                    except Exception as exc:  # noqa: BLE001
                        had_error = True
                        logger.warning("DOCX extract failed (%s): %s", f.name, exc)
                elif not unsupported:
                    unsupported = True  # unknown extension → flag, do not guess

                word_count = len(text.split())
                status = _classify(word_count, had_error, unsupported)

                entries.append(Entry(
                    anon_id=anon,
                    cohort=cohort,
                    journal_index=journal_index,
                    canvas_user_id=parsed["canvas_id"],
                    submission_id=parsed["submission_id"],
                    is_late=parsed["is_late"],
                    is_resubmission=parsed["is_resubmission"],
                    resubmission_n=parsed["resubmission_n"],
                    file_ext=ext,
                    file_bytes=f.stat().st_size,
                    extract_status=status,
                    word_count=word_count,
                    char_count=len(text),
                    doc_created=created,
                    doc_modified=modified,
                    source_path=str(f),
                    text=text if status in (STATUS_OK, STATUS_SUSPECT) else "",
                ))

    df = pd.DataFrame([asdict(e) for e in entries])
    _write_outputs(df, unparsed, crosswalk)
    _print_summary(df, unparsed, crosswalk)
    return df


def _write_outputs(
    df: pd.DataFrame,
    unparsed: list[dict],
    crosswalk: dict[tuple[str, str], set[str]],
) -> None:
    PROCESSED_DIR.mkdir(parents=True, exist_ok=True)
    CROSSWALK_DIR.mkdir(parents=True, exist_ok=True)

    df.to_parquet(ENTRIES_PARQUET, index=False)
    df.drop(columns=["text"]).to_csv(ENTRY_MANIFEST_CSV, index=False)

    pd.DataFrame(unparsed).to_csv(UNPARSED_CSV, index=False)

    rows = []
    for (cohort, name), ids in sorted(crosswalk.items()):
        rows.append({
            "anon_id": anon_id_for(cohort, name),
            "cohort": cohort,
            "normalised_name": name,
            "canvas_user_ids": "|".join(sorted(ids)),
            "id_collision": len(ids) > 1,
        })
    pd.DataFrame(rows).to_csv(CROSSWALK_CSV, index=False)


def _print_summary(
    df: pd.DataFrame,
    unparsed: list[dict],
    crosswalk: dict[tuple[str, str], set[str]],
) -> None:
    print(f"Ingested {len(df)} journal files across "
          f"{df['cohort'].nunique()} cohorts "
          f"({len(crosswalk)} distinct students).")
    print(f"  unparsed filenames : {len(unparsed)}")
    print("  extract_status     :")
    for status, n in df["extract_status"].value_counts().items():
        print(f"      {status:<18} {n}")
    print(f"  late submissions   : {int(df['is_late'].sum())}")
    print(f"  resubmissions      : {int(df['is_resubmission'].sum())}")
    dates = (df["doc_created"] != "").sum()
    print(f"  doc_created present: {dates}/{len(df)} "
          f"({100 * dates / max(len(df), 1):.0f}%)")
    print(f"Wrote: {ENTRIES_PARQUET}, {ENTRY_MANIFEST_CSV}, {CROSSWALK_CSV}")


def main(argv: list[str] | None = None) -> None:
    parser = argparse.ArgumentParser(
        prog="python3 -m src.qualitative.ingest",
        description="Ingest Canvas journal exports into a pseudonymised table.",
    )
    parser.add_argument(
        "--raw-dir", type=Path, default=RAW_DIR,
        help=f"Raw export root (default: {RAW_DIR}).",
    )
    args = parser.parse_args(argv)
    logging.basicConfig(level=logging.INFO, format="%(levelname)s %(message)s")
    ingest(args.raw_dir)


if __name__ == "__main__":
    main()
